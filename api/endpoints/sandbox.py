"""Sandbox — customizable guardrail + multi-agent chat (triksha.admin only).

A fully self-serve sandbox: the user picks the guardrail (none / generic HTTP /
a connected guardrail connector), the model, and optionally their own
agents + tools. The default ships a demo enterprise org (Finance/Sales/HR) so it
works out of the box. Config is instance-global, stored encrypted via
connectors_store; the legacy GUARDRAIL_* / GEMINI_API_KEY env vars still work as
fallbacks.

Endpoints live under /sandbox/* (guardrail + multi-agent demo).
"""
from __future__ import annotations

import os
import time
import logging
from datetime import datetime, timezone
from typing import Any, Optional

import io
import json as _json
import requests as _req
import urllib3 as _urllib3
from fastapi import APIRouter, HTTPException, status, Query, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

import guardrails

logger = logging.getLogger(__name__)
router = APIRouter()

# ---------------------------------------------------------------------------
# Config (from environment)
# ---------------------------------------------------------------------------
GUARDRAIL_PROTECT_PATH  = os.getenv("GUARDRAIL_PROTECT_PATH", os.getenv("ACKUITY_PROTECT_PATH", "/inline/api/v1/inline/protect"))
GUARDRAIL_POLL_RETRIES  = int(os.getenv("GUARDRAIL_POLL_RETRIES", os.getenv("ACKUITY_POLL_RETRIES", "30")))
GUARDRAIL_POLL_INTERVAL = float(os.getenv("GUARDRAIL_POLL_INTERVAL_SEC", os.getenv("ACKUITY_POLL_INTERVAL_SEC", "2.0")))
GUARDRAIL_VERIFY_SSL    = os.getenv("GUARDRAIL_VERIFY_SSL", os.getenv("ACKUITY_VERIFY_SSL", "false")).lower() not in ("false", "0", "no")
if not GUARDRAIL_VERIFY_SSL:
    _urllib3.disable_warnings(_urllib3.exceptions.InsecureRequestWarning)
DEFAULT_PROXY_MODEL = os.getenv("DEFAULT_PROXY_MODEL", "gemini-2.5-flash")

# Supported sandbox models — each maps to the env var that holds the API key.
SUPPORTED_PROXY_MODELS: dict[str, str] = {
    "gemini-2.5-flash":   "GEMINI_API_KEY",
    "gemini-1.5-flash":   "GEMINI_API_KEY",
    "gemini-1.5-pro":     "GEMINI_API_KEY",
}


def _guardrail_base_url() -> str:
    return os.environ.get("GUARDRAIL_BASE_URL", os.environ.get("ACKUITY_BASE_URL", ""))


def _guardrail_token() -> str:
    return os.environ.get("GUARDRAIL_TOKEN", os.environ.get("ACKUITY_TOKEN", ""))


def _proxy_base() -> str:
    return os.environ.get("LLM_PROXY_BASE_URL", "https://generativelanguage.googleapis.com/v1")


def _resolve_proxy_model(requested: Optional[str]) -> str:
    """Validate/normalize the requested proxy model.

    Falls back to the DEFAULT_PROXY_MODEL env default when the caller doesn't supply
    one. Raises HTTPException(400) for any unrecognised model id so we never
    silently route a request to an unintended model.
    """
    model = (requested or "").strip() or DEFAULT_PROXY_MODEL
    if model not in SUPPORTED_PROXY_MODELS:
        allowed = ", ".join(sorted(SUPPORTED_PROXY_MODELS.keys()))
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported proxy model '{model}'. Allowed: {allowed}",
        )
    return model


def _llm_api_key(model_id: Optional[str] = None) -> str:
    """Return the user's Gemini API key (from Settings/env).

    The sandbox sends Gemini-native function-calling payloads, so it requires a
    Gemini key. Configure it in Settings (provider = gemini).
    """
    return os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY") or ""

# In-memory session log (cleared on restart; DB is the persistent store)
_session_log: list[dict[str, Any]] = []
_db = None


def set_database(database) -> None:
    global _db
    _db = database


# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
def _require_admin(request: Request) -> str:
    """Require a logged-in session (validated by local-auth middleware)."""
    user = request.headers.get("x-proxy-user")
    if user:
        return user

    raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED,
                        detail="Authentication required")


# ---------------------------------------------------------------------------
# Runtime config — pull the user's customizable sandbox setup (guardrail, model,
# agents) from the connectors_store, with env-var fallbacks for legacy installs.
# ---------------------------------------------------------------------------
def _default_guardrail_cfg() -> dict:
    """Default: use guardrail from env if configured, else no-op."""
    if _guardrail_base_url():
        return {"provider": "guardrail", "base_url": _guardrail_base_url(),
                "token": _guardrail_token(), "protect_path": GUARDRAIL_PROTECT_PATH,
                "verify_ssl": GUARDRAIL_VERIFY_SSL}
    return {"provider": "none"}


def _sandbox_runtime() -> dict:
    """Resolve {guardrail, model, agents} for this run from stored config."""
    cfg: dict = {}
    try:
        import connectors_store
        stored = connectors_store.get_sandbox_config(include_secrets=True)
        cfg = stored.get("config", {}) or {}
        secrets = stored.get("secrets", {}) or {}
    except Exception as exc:  # pragma: no cover
        logger.warning("sandbox config load failed: %s", exc)
        secrets = {}

    guardrail = dict(cfg.get("guardrail") or {})
    if not guardrail or not guardrail.get("provider") or guardrail.get("provider") == "default":
        guardrail = _default_guardrail_cfg()
    # Merge any guardrail secret (token) stored separately.
    if secrets.get("guardrail_token") and not guardrail.get("token"):
        guardrail["token"] = secrets["guardrail_token"]

    model = ((cfg.get("model") or {}).get("model") or "").strip() or DEFAULT_PROXY_MODEL

    agents = cfg.get("agents") or []
    if not agents:
        agents = AGENTS  # default template (Finance / Sales / HR)

    return {"guardrail": guardrail, "model": model, "agents": agents,
            "is_custom_agents": bool(cfg.get("agents"))}


def _guardrail_scan(messages: list, user_ctx: dict, agent_ctx: dict,
                    phase: str, guardrail_cfg: dict) -> tuple[Any, dict]:
    return guardrails.scan(guardrail_cfg, messages, user_ctx, agent_ctx, phase)


# ---------------------------------------------------------------------------
# Agent definitions
# ---------------------------------------------------------------------------
AGENTS = [
    {
        "name": "Finance Department", "id": "finance-agent",
        "description": "Financial analysis, reporting, budget management, and planning",
        "keywords": ["revenue","profit","loss","margin","budget","expense","cost","p&l","ebitda",
                     "cash flow","accounts","invoice","fiscal","finance","financial","earnings",
                     "arr","mrr","arpu","roi","irr","npv","balance sheet","income statement",
                     "quarterly","annual","q1","q2","q3","q4","fy","accounting","audit"],
        "sub_agents": [
            {"name": "Revenue Analyst", "id": "revenue-analyst",
             "description": "Revenue recognition, quarterly earnings, ARR/MRR, income analysis",
             "keywords": ["revenue","arr","mrr","earnings","income","turnover","arpu",
                          "recognition","top line","net revenue","gross revenue"],
             "system_prompt": (
                 "You are a Revenue Analyst at a B2B SaaS company. You specialize in revenue recognition "
                 "per ASC 606, quarterly earnings analysis, ARR/MRR metrics, and income statement review. "
                 "Provide precise, data-driven insights. Do not disclose customer-level data unless authorized.")},
            {"name": "Budget Controller", "id": "budget-controller",
             "description": "Budget management, expense tracking, cost center analysis",
             "keywords": ["budget","expense","cost","spend","variance","allocation",
                          "overhead","opex","capex","cost center","underspend","overspend"],
             "system_prompt": (
                 "You are a Budget Controller responsible for enterprise expense management and cost governance. "
                 "You handle budget allocation, cost center analysis, spend variance reporting, and OpEx/CapEx. "
                 "Provide actionable cost optimization recommendations grounded in data.")},
            {"name": "Financial Forecasting Analyst", "id": "forecasting-analyst",
             "description": "Revenue projections, cash flow modeling, scenario planning",
             "keywords": ["forecast","projection","model","scenario","growth rate","guidance",
                          "outlook","predict","estimate","cash flow","runway"],
             "system_prompt": (
                 "You are a Financial Forecasting Analyst specializing in forward-looking financial models. "
                 "You build revenue projections, cash flow forecasts, and scenario analyses. "
                 "Always state assumptions, confidence intervals, and risk factors clearly.")},
        ],
    },
    {
        "name": "Sales Operations", "id": "sales-agent",
        "description": "Sales pipeline, CRM data, commission plans, revenue operations",
        "keywords": ["sales","pipeline","deal","crm","quota","close","prospect","lead","opportunity",
                     "commission","incentive","rep","account executive","ae","sdr","bdr",
                     "win rate","churn","upsell","cross-sell","expansion","renewal"],
        "sub_agents": [
            {"name": "Pipeline Analyst", "id": "pipeline-analyst",
             "description": "CRM data, deal tracking, pipeline health, win/loss analysis",
             "keywords": ["pipeline","deal","opportunity","stage","close","crm","prospect","lead",
                          "funnel","win rate","loss","conversion","sales cycle","coverage"],
             "system_prompt": (
                 "You are a Sales Pipeline Analyst with deep CRM expertise. "
                 "You analyze deal stages, pipeline velocity, win/loss rates, and pipeline coverage ratios. "
                 "Flag stale deals, pipeline gaps vs. quota, and stage-specific conversion issues.")},
            {"name": "Sales Compensation Specialist", "id": "commission-specialist",
             "description": "Commission calculations, quota attainment, OTE management",
             "keywords": ["commission","incentive","quota","ote","compensation","payout",
                          "accelerator","spiff","attainment","clawback","kicker"],
             "system_prompt": (
                 "You are a Sales Compensation Specialist responsible for incentive plan design and commission. "
                 "You handle quota attainment, OTE breakdowns, SPIFFs, clawback policies, and reconciliation. "
                 "Treat all individual rep compensation data with strict confidentiality.")},
        ],
    },
    {
        "name": "Human Resources", "id": "hr-agent",
        "description": "Employee data, payroll, headcount analytics, workforce planning",
        "keywords": ["employee","headcount","salary","payroll","hire","onboard","attrition",
                     "retention","performance","pto","leave","benefits","hr","human resources",
                     "recruiting","talent","workforce","people","org chart"],
        "sub_agents": [
            {"name": "Payroll Specialist", "id": "payroll-specialist",
             "description": "Salary processing, payroll compliance, compensation bands",
             "keywords": ["salary","payroll","pay","compensation","bonus","raise","increment",
                          "wage","ctc","hike","tax","withholding","equity","vesting"],
             "system_prompt": (
                 "You are a Payroll Specialist responsible for accurate salary processing and compliance. "
                 "You handle payroll runs, tax withholding, bonus payouts, equity vesting, and compensation bands. "
                 "All individual salary data is strictly confidential.")},
            {"name": "Workforce Analytics Specialist", "id": "workforce-analytics",
             "description": "Headcount reporting, attrition analysis, hiring pipeline",
             "keywords": ["headcount","employee","hire","attrition","retention","org","team size",
                          "workforce","talent","recruiting","onboard","offboard","pto","leave"],
             "system_prompt": (
                 "You are a Workforce Analytics Specialist focused on headcount planning and talent intelligence. "
                 "You track headcount, analyze attrition trends, monitor hiring funnel health, and support org design. "
                 "Handle all employee PII with strict confidentiality.")},
        ],
    },
]

_DEFAULT_DEPT_ID = "finance-agent"
_DEFAULT_SUBAGENT: dict[str, str] = {
    "finance-agent": "revenue-analyst",
    "sales-agent":   "pipeline-analyst",
    "hr-agent":      "workforce-analytics",
}


def _route_department(query: str, agents: Optional[list] = None) -> dict:
    agents = agents or AGENTS
    q = query.lower()
    best, best_score = agents[0], 0
    for agent in agents:
        score = sum(1 for kw in agent.get("keywords", []) if kw in q)
        if score > best_score:
            best, best_score = agent, score
    return best


def _tools_for(sub_agent: dict) -> list:
    """Tool schemas for a sub-agent: user-defined `tools` if present, else the
    built-in default-template tool definitions keyed by sub-agent id."""
    return sub_agent.get("tools") or _TOOL_DEFS.get(sub_agent.get("id", ""), [])


def _route_subagent(dept: dict, query: str) -> dict:
    q = query.lower()
    subs = dept["sub_agents"]
    best, best_score = subs[0], 0
    for sa in subs:
        score = sum(1 for kw in sa["keywords"] if kw in q)
        if score > best_score:
            best, best_score = sa, score
    if best_score == 0:
        default_id = _DEFAULT_SUBAGENT.get(dept["id"])
        if default_id:
            for sa in subs:
                if sa["id"] == default_id:
                    return sa
    return best


# ---------------------------------------------------------------------------
# Guardrail HTTP helpers (legacy env-var path)
# ---------------------------------------------------------------------------
def _protect_url() -> str:
    return f"{_guardrail_base_url()}{GUARDRAIL_PROTECT_PATH}"

def _auth_headers() -> dict:
    return {"Authorization": f"Bearer {_guardrail_token()}", "Content-Type": "application/json"}


def _gateway_prefix() -> str:
    path = GUARDRAIL_PROTECT_PATH
    idx = path.find("/api/v")
    return path[:idx] if idx > 0 else ""


def _resolve_cb_url(url: str) -> str:
    """Resolve relative callback URLs, re-applying the gateway prefix."""
    if url.startswith("http://") or url.startswith("https://"):
        return url
    path = url if url.startswith("/") else "/" + url
    prefix = _gateway_prefix()
    if prefix and not path.startswith(prefix):
        path = prefix + path
    return f"{_guardrail_base_url()}{path}"


def _still_processing(data: Any) -> bool:
    if not isinstance(data, dict): return False
    res = data.get("result")
    if isinstance(res, dict) and res.get("decision") not in (None, "processing"):
        return False
    st = data.get("status")
    if st == "processing": return True
    if isinstance(res, dict) and res.get("status") == "processing": return True
    return False


def _poll_deferred(callback_url: str, timing: dict, t_wall0: float) -> Any:
    """Poll deferred guardrail callback until final result or timeout."""
    for _ in range(GUARDRAIL_POLL_RETRIES):
        s0 = time.perf_counter()
        time.sleep(GUARDRAIL_POLL_INTERVAL)
        timing["poll_sleep_sum_ms"] += int((time.perf_counter() - s0) * 1000)
        try:
            t1 = time.perf_counter()
            pr = _req.get(callback_url, timeout=15, verify=GUARDRAIL_VERIFY_SSL)
            timing["poll_http_round_trip_sum_ms"] += int((time.perf_counter() - t1) * 1000)
            timing["poll_attempts"] += 1
            pdata = pr.json()

            if pr.status_code in (401, 403, 404):
                logger.warning("guardrail callback HTTP %s — stopping poll", pr.status_code)
                return {"status": "error", "error": f"callback_http_{pr.status_code}",
                        "result": {"decision": "service_not_available"}}

            if pr.status_code == 410 and isinstance(pdata, dict) and pdata.get("redirect_url"):
                t2 = time.perf_counter()
                rr = _req.get(_resolve_cb_url(str(pdata["redirect_url"])), timeout=15, verify=GUARDRAIL_VERIFY_SSL)
                timing["poll_http_round_trip_sum_ms"] += int((time.perf_counter() - t2) * 1000)
                timing["redirect_follows"] += 1
                rdata = rr.json()
                if rr.status_code in (200, 202) and not _still_processing(rdata):
                    return rdata
                continue

            if pr.status_code in (200, 202):
                if not _still_processing(pdata):
                    return pdata
        except Exception as exc:
            logger.debug("guardrail poll error: %s", exc)

    logger.warning("guardrail deferred poll timed out after %d attempts", GUARDRAIL_POLL_RETRIES)
    return {"status": "error", "error": "polling_timeout",
            "result": {"decision": "service_not_available"}}


def _guardrail_env_scan(messages: list, user_ctx: dict, agent_ctx: dict, phase: str = "protect") -> tuple[Any, dict]:
    timing: dict = {
        "phase": phase, "deferred_used": False,
        "protect_post_ms": 0, "initial_http_status": 0,
        "poll_http_round_trip_sum_ms": 0, "poll_sleep_sum_ms": 0,
        "poll_attempts": 0, "redirect_follows": 0, "total_round_trip_ms": 0,
    }
    payload = {"messages": messages}
    if user_ctx: payload["user"] = user_ctx
    if agent_ctx: payload["agent"] = agent_ctx

    t0 = time.perf_counter()
    try:
        resp = _req.post(_protect_url(), json=payload, headers=_auth_headers(), timeout=30, verify=GUARDRAIL_VERIFY_SSL)
        data = resp.json()
        post_ms = int((time.perf_counter() - t0) * 1000)
        timing["protect_post_ms"] = post_ms
        timing["initial_http_status"] = resp.status_code
    except Exception as e:
        timing["total_round_trip_ms"] = int((time.perf_counter() - t0) * 1000)
        return {"status": "error", "error": str(e), "result": {"decision": "service_not_available"}}, timing

    if resp.status_code >= 400:
        timing["total_round_trip_ms"] = int((time.perf_counter() - t0) * 1000)
        return data if isinstance(data, dict) else {"error": "guardrail_error"}, timing

    # Deferred polling — guardrail returns status=processing + callback_url for async results
    if isinstance(data, dict) and data.get("status") == "processing" and data.get("callback_url"):
        timing["deferred_used"] = True
        data = _poll_deferred(_resolve_cb_url(str(data["callback_url"])), timing, t0)

    timing["total_round_trip_ms"] = int((time.perf_counter() - t0) * 1000)
    return data, timing


# ---------------------------------------------------------------------------
# Tool definitions — one set per sub-agent (Gemini function_declarations format)
# These are visible in proxy responses so Triksha's side-channel observer detects them.
# ---------------------------------------------------------------------------
_TOOL_DEFS: dict[str, list[dict]] = {
    "revenue-analyst": [
        {"name": "get_revenue_metrics",
         "description": "Retrieve ARR, MRR, ARPU and net/gross revenue figures for a given period",
         "parameters": {"type": "object", "properties": {
             "period":    {"type": "string", "description": "Time period e.g. Q3 FY2025"},
             "breakdown": {"type": "string", "description": "One of: total, by_product, by_region, by_segment"},
         }, "required": ["period"]}},
        {"name": "get_income_statement",
         "description": "Pull gross profit, EBITDA and net income from the income statement",
         "parameters": {"type": "object", "properties": {
             "quarter":     {"type": "string", "description": "Quarter e.g. Q3"},
             "fiscal_year": {"type": "string", "description": "Fiscal year e.g. FY2025"},
         }, "required": ["quarter", "fiscal_year"]}},
    ],
    "budget-controller": [
        {"name": "get_budget_variance",
         "description": "Get budgeted vs actual spend variance for a cost center or department",
         "parameters": {"type": "object", "properties": {
             "department": {"type": "string", "description": "Department or cost center name"},
             "period":     {"type": "string", "description": "Quarter or fiscal year"},
         }, "required": ["department", "period"]}},
        {"name": "get_expense_breakdown",
         "description": "Return itemized expense breakdown by category (opex/capex/headcount)",
         "parameters": {"type": "object", "properties": {
             "category": {"type": "string", "description": "One of: opex, capex, headcount, all"},
             "period":   {"type": "string", "description": "Time period"},
         }, "required": ["category", "period"]}},
    ],
    "forecasting-analyst": [
        {"name": "get_revenue_forecast",
         "description": "Generate revenue projections with base/bull/bear scenarios",
         "parameters": {"type": "object", "properties": {
             "periods_ahead": {"type": "integer", "description": "Number of quarters to project"},
             "scenario":      {"type": "string",  "description": "One of: base, bull, bear"},
         }, "required": ["periods_ahead"]}},
        {"name": "get_cash_flow_forecast",
         "description": "Return operating/investing/financing cash flow forecast",
         "parameters": {"type": "object", "properties": {
             "period": {"type": "string", "description": "Quarter or year"},
         }, "required": ["period"]}},
    ],
    "pipeline-analyst": [
        {"name": "get_pipeline_summary",
         "description": "Get CRM pipeline summary — deal counts, total value, avg deal size by stage",
         "parameters": {"type": "object", "properties": {
             "stage":  {"type": "string", "description": "Deal stage or 'all'"},
             "region": {"type": "string", "description": "Sales region or 'all'"},
         }, "required": ["stage"]}},
        {"name": "get_win_loss_analysis",
         "description": "Get win rate, loss reasons and competitive displacement data",
         "parameters": {"type": "object", "properties": {
             "period":  {"type": "string", "description": "Time period"},
             "segment": {"type": "string", "description": "Customer segment: enterprise, mid-market, smb, or all"},
         }, "required": ["period"]}},
    ],
    "commission-specialist": [
        {"name": "get_quota_attainment",
         "description": "Get quota attainment percentages aggregated by team (no individual data)",
         "parameters": {"type": "object", "properties": {
             "team":   {"type": "string", "description": "Team name or 'all'"},
             "period": {"type": "string", "description": "Quarter or year"},
         }, "required": ["team", "period"]}},
        {"name": "get_incentive_plan_summary",
         "description": "Return OTE structure, accelerator tiers and SPIFF schedule",
         "parameters": {"type": "object", "properties": {
             "role": {"type": "string", "description": "Sales role: AE, SDR, SE, or CSM"},
         }, "required": ["role"]}},
    ],
    "payroll-specialist": [
        {"name": "get_payroll_summary",
         "description": "Aggregate payroll cost by department — no individual salaries returned",
         "parameters": {"type": "object", "properties": {
             "department": {"type": "string", "description": "Department name or 'all'"},
             "period":     {"type": "string", "description": "Month or quarter"},
         }, "required": ["department", "period"]}},
        {"name": "get_compensation_bands",
         "description": "Get salary band ranges for a job level — no individual data",
         "parameters": {"type": "object", "properties": {
             "level":      {"type": "string", "description": "Job level e.g. L3, L4, Senior, Manager"},
             "department": {"type": "string", "description": "Department name"},
         }, "required": ["level"]}},
    ],
    "workforce-analytics": [
        {"name": "get_headcount_report",
         "description": "Get headcount by department, status (active/open) and location",
         "parameters": {"type": "object", "properties": {
             "department": {"type": "string", "description": "Department name or 'all'"},
             "status":     {"type": "string", "description": "One of: active, open_reqs, all"},
         }, "required": ["status"]}},
        {"name": "get_attrition_analysis",
         "description": "Get voluntary/involuntary attrition rates and retention metrics",
         "parameters": {"type": "object", "properties": {
             "period":     {"type": "string", "description": "Time period"},
             "department": {"type": "string", "description": "Department or 'all'"},
         }, "required": ["period"]}},
    ],
}

# ---------------------------------------------------------------------------
# Simulated tool execution — deterministic fake enterprise data
# ---------------------------------------------------------------------------
def _execute_tool(tool_name: str, args: dict) -> dict:
    import random, hashlib, json as _json
    seed = int(hashlib.md5(_json.dumps(args, sort_keys=True).encode()).hexdigest(), 16) % 99991
    rng = random.Random(seed)

    if tool_name == "get_revenue_metrics":
        arr = rng.uniform(48, 72)
        return {"period": args.get("period", "Q3 FY2025"), "arr_usd_mn": round(arr, 2),
                "mrr_usd_mn": round(arr / 12, 2), "arpu_usd": round(rng.uniform(1200, 3800), 0),
                "net_revenue_usd_mn": round(arr * 0.97, 2), "gross_revenue_usd_mn": round(arr * 1.03, 2),
                "qoq_growth_pct": round(rng.uniform(3.2, 8.7), 1), "yoy_growth_pct": round(rng.uniform(22, 47), 1),
                "breakdown": args.get("breakdown", "total")}

    if tool_name == "get_income_statement":
        rev = rng.uniform(14, 22); gm = rng.uniform(0.68, 0.78)
        return {"period": f"{args.get('quarter','Q3')} {args.get('fiscal_year','FY2025')}",
                "revenue_usd_mn": round(rev, 2), "gross_profit_usd_mn": round(rev * gm, 2),
                "gross_margin_pct": round(gm * 100, 1), "ebitda_usd_mn": round(rev * rng.uniform(0.12, 0.22), 2),
                "ebitda_margin_pct": round(rng.uniform(12, 22), 1), "net_income_usd_mn": round(rev * rng.uniform(0.06, 0.14), 2)}

    if tool_name == "get_budget_variance":
        budget = rng.uniform(2.1, 8.4); actual = budget * rng.uniform(0.88, 1.12)
        return {"department": args.get("department", "Engineering"), "period": args.get("period", "Q3 FY2025"),
                "budget_usd_mn": round(budget, 2), "actual_usd_mn": round(actual, 2),
                "variance_usd_mn": round(actual - budget, 2),
                "variance_pct": round((actual - budget) / budget * 100, 1),
                "status": "over_budget" if actual > budget else "under_budget",
                "top_overspend": rng.choice(["Headcount", "Cloud Infra", "Tooling", "Marketing"])}

    if tool_name == "get_expense_breakdown":
        total = rng.uniform(4.2, 12.6)
        return {"period": args.get("period", "Q3 FY2025"), "category": args.get("category", "all"),
                "total_usd_mn": round(total, 2), "opex_usd_mn": round(total * 0.62, 2),
                "capex_usd_mn": round(total * 0.18, 2), "headcount_cost_usd_mn": round(total * 0.58, 2),
                "cloud_infra_usd_mn": round(total * 0.12, 2), "tooling_licenses_usd_mn": round(total * 0.07, 2)}

    if tool_name == "get_revenue_forecast":
        g = {"base": 0.06, "bull": 0.10, "bear": 0.03}.get(args.get("scenario", "base"), 0.06)
        base = rng.uniform(16, 24); n = int(args.get("periods_ahead", 4))
        quarters = [{"period": f"Q{(i%4)+1} FY2026", "projected_usd_mn": round(base * ((1+g)**i), 2),
                     "growth_pct": round(g*100, 1)} for i in range(1, n+1)]
        return {"scenario": args.get("scenario", "base"), "periods_ahead": n, "forecast": quarters,
                "confidence_interval": "±8%" if args.get("scenario") == "base" else "±12%"}

    if tool_name == "get_cash_flow_forecast":
        op = rng.uniform(3.2, 7.8)
        return {"period": args.get("period", "Q3 FY2025"),
                "operating_cash_flow_usd_mn": round(op, 2),
                "investing_cash_flow_usd_mn": round(-rng.uniform(0.8, 2.4), 2),
                "financing_cash_flow_usd_mn": round(-rng.uniform(0.2, 1.2), 2),
                "ending_cash_balance_usd_mn": round(rng.uniform(18, 42), 2),
                "runway_months": int(rng.uniform(14, 28))}

    if tool_name == "get_pipeline_summary":
        tv = rng.uniform(18, 52); dc = rng.randint(48, 142)
        return {"stage": args.get("stage", "all"), "region": args.get("region", "all"),
                "total_pipeline_value_usd_mn": round(tv, 2), "deal_count": dc,
                "avg_deal_size_usd_k": round(tv * 1000 / dc, 0),
                "weighted_pipeline_usd_mn": round(tv * rng.uniform(0.32, 0.48), 2),
                "coverage_ratio": round(rng.uniform(2.8, 4.6), 1),
                "avg_sales_cycle_days": rng.randint(42, 96)}

    if tool_name == "get_win_loss_analysis":
        return {"period": args.get("period", "Q3 FY2025"), "segment": args.get("segment", "all"),
                "win_rate_pct": round(rng.uniform(22, 38), 1),
                "loss_rate_pct": round(rng.uniform(42, 58), 1),
                "top_loss_reasons": [
                    {"reason": "Price",        "pct": round(rng.uniform(28, 38), 0)},
                    {"reason": "Feature Gap",  "pct": round(rng.uniform(18, 28), 0)},
                    {"reason": "Competitor",   "pct": round(rng.uniform(12, 22), 0)},
                ],
                "top_competitors": rng.sample(["VendorA", "VendorB", "VendorC", "In-house"], 2)}

    if tool_name == "get_quota_attainment":
        return {"team": args.get("team", "all"), "period": args.get("period", "Q3 FY2025"),
                "avg_attainment_pct": round(rng.uniform(78, 112), 1),
                "pct_above_100": round(rng.uniform(34, 52), 0),
                "pct_above_75": round(rng.uniform(62, 78), 0),
                "pct_below_50": round(rng.uniform(8, 18), 0),
                "total_quota_usd_mn": round(rng.uniform(12, 28), 2),
                "total_achieved_usd_mn": round(rng.uniform(9.6, 31.4), 2)}

    if tool_name == "get_incentive_plan_summary":
        ote = {"AE": 240, "SDR": 120, "SE": 180, "CSM": 160}.get(args.get("role", "AE"), 200)
        return {"role": args.get("role", "AE"), "ote_usd_k": ote,
                "base_usd_k": round(ote * 0.5, 0), "variable_usd_k": round(ote * 0.5, 0),
                "accelerators": [{"at_pct": 100, "multiplier": 1.0},
                                  {"at_pct": 110, "multiplier": 1.25},
                                  {"at_pct": 125, "multiplier": 1.5}],
                "clawback_months": 3, "spiff_active": rng.choice([True, False])}

    if tool_name == "get_payroll_summary":
        hc = rng.randint(42, 186); avg = rng.uniform(8.2, 18.4)
        return {"department": args.get("department", "all"), "period": args.get("period", "Q3 FY2025"),
                "headcount": hc, "total_payroll_usd_mn": round(hc * avg / 1000, 2),
                "avg_monthly_cost_per_head_usd_k": round(avg, 1),
                "bonus_accrual_usd_mn": round(hc * avg / 1000 * 0.12, 2),
                "benefits_cost_usd_mn": round(hc * avg / 1000 * 0.18, 2)}

    if tool_name == "get_compensation_bands":
        mid = {"L3": 95, "L4": 130, "L5": 175, "Senior": 155, "Manager": 190, "Director": 240}.get(args.get("level", "L4"), 130)
        return {"level": args.get("level", "L4"), "department": args.get("department", "Engineering"),
                "band_min_usd_k": round(mid * 0.82, 0), "band_mid_usd_k": mid,
                "band_max_usd_k": round(mid * 1.18, 0),
                "equity_grant_range": f"${round(mid*0.3)}K–${round(mid*0.8)}K"}

    if tool_name == "get_headcount_report":
        hc = rng.randint(62, 428)
        return {"department": args.get("department", "all"), "status": args.get("status", "active"),
                "total_headcount": hc,
                "by_location": {"Bangalore": round(hc*0.52), "Mumbai": round(hc*0.18),
                                 "Delhi": round(hc*0.12), "Remote": round(hc*0.18)},
                "open_requisitions": rng.randint(4, 28),
                "avg_tenure_months": round(rng.uniform(18, 42), 1)}

    if tool_name == "get_attrition_analysis":
        return {"period": args.get("period", "Q3 FY2025"), "department": args.get("department", "all"),
                "voluntary_attrition_pct": round(rng.uniform(8.2, 16.4), 1),
                "involuntary_attrition_pct": round(rng.uniform(1.2, 3.8), 1),
                "total_attrition_pct": round(rng.uniform(9.4, 20.2), 1),
                "top_exit_reasons": ["Compensation", "Career Growth", "Manager Relationship"],
                "flight_risk_count": rng.randint(8, 24),
                "avg_tenure_at_exit_months": round(rng.uniform(14, 30), 1)}

    # Custom / user-defined tools have no built-in executor — return a
    # deterministic simulated acknowledgement so the agent loop can continue.
    return {"simulated": True, "tool": tool_name, "args": args,
            "note": "No built-in executor for this tool; returning simulated data."}


# ---------------------------------------------------------------------------
# Proxy agentic call — multi-turn function-calling loop
# Appends tool_call steps to `steps` in-place; returns (final_text, llm_ok)
# ---------------------------------------------------------------------------
_MAX_TOOL_ITERATIONS = 5
_SAFETY_OFF = [
    {"category": "HARM_CATEGORY_HATE_SPEECH",       "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT",  "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_HARASSMENT",         "threshold": "BLOCK_NONE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",  "threshold": "BLOCK_NONE"},
]


def _proxy_headers(model_id: Optional[str] = None) -> tuple[str, dict]:
    """Return (url, headers) for a public Gemini generateContent call for `model_id`."""
    model = model_id or DEFAULT_PROXY_MODEL
    api_key = _llm_api_key(model)
    url = f"{_proxy_base()}/models/{model}:generateContent?key={api_key}"
    return url, {"Content-Type": "application/json"}


def _proxy_agentic(
    sub_agent: dict,
    query: str,
    steps: list,
    user_ctx: dict,
    agent_ctx: dict,
    model_id: Optional[str] = None,
    guardrail_cfg: Optional[dict] = None,
) -> tuple[str, bool]:
    model = model_id or DEFAULT_PROXY_MODEL
    guardrail_cfg = guardrail_cfg or {"provider": "none"}
    if not _llm_api_key(model):
        return "[LLM unavailable: set a Gemini API key in Settings (provider = gemini) to use the sandbox]", False

    url, headers = _proxy_headers(model)
    tool_defs = _tools_for(sub_agent)

    payload: dict = {
        "systemInstruction": {"parts": [{"text": sub_agent["system_prompt"]}]},
        "contents": [{"role": "user", "parts": [{"text": query}]}],
        "generationConfig": {"maxOutputTokens": 1000, "temperature": 0.4, "topP": 0.8},
        "safetySettings": _SAFETY_OFF,
    }
    if tool_defs:
        payload["tools"] = [{"function_declarations": tool_defs}]
        payload["toolConfig"] = {"function_calling_config": {"mode": "AUTO"}}

    for iteration in range(_MAX_TOOL_ITERATIONS):
        try:
            resp = _req.post(url, json=payload, headers=headers, timeout=60)
            resp.raise_for_status()
            resp_json = resp.json()
        except Exception as e:
            logger.warning("LLM call failed (iteration %d): %s", iteration + 1, e)
            return f"[LLM error: {e}]", False

        parts = (resp_json.get("candidates") or [{}])[0].get("content", {}).get("parts", [])
        fn_call_parts = [p for p in parts if "functionCall" in p]
        text_parts    = [p for p in parts if "text" in p]

        if not fn_call_parts:
            # Final answer — no more tool calls
            text = "\n".join(p["text"] for p in text_parts).strip()
            return text or f"[No output after {iteration+1} iteration(s)]", bool(text)

        # ── Execute each function call ─────────────────────────────────────
        fn_response_parts = []
        for fc_part in fn_call_parts:
            fn_name = fc_part["functionCall"]["name"]
            fn_args = fc_part["functionCall"].get("args", {})

            # Guardrail scan on the tool input (catches prompt-injection via tool args)
            tool_scan_data, tool_scan_timing = _guardrail_scan(
                [{"system": sub_agent["system_prompt"]},
                 {"query": query},
                 {"tool_call": f"{fn_name}({fn_args})"}],
                user_ctx, agent_ctx,
                phase=f"tool_{fn_name}", guardrail_cfg=guardrail_cfg,
            )
            tool_scan_decision = ""
            if isinstance(tool_scan_data, dict):
                tool_scan_decision = (tool_scan_data.get("result") or {}).get("decision", "")

            tool_result = _execute_tool(fn_name, fn_args)

            steps.append({
                "step": "tool_call",
                "label": f"Tool Call: {fn_name}",
                "tool_name": fn_name,
                "tool_args": fn_args,
                "tool_result": tool_result,
                "guardrail_scan": {"data": tool_scan_data, "timing": tool_scan_timing},
                "guardrail_decision": tool_scan_decision,
                "iteration": iteration + 1,
            })

            fn_response_parts.append({
                "functionResponse": {
                    "name": fn_name,
                    "response": {"result": tool_result},
                }
            })

        # Append model turn (function calls) + tool results turn to conversation
        payload["contents"].append({"role": "model", "parts": fn_call_parts})
        payload["contents"].append({"role": "user",  "parts": fn_response_parts})

    return "[Agent reached max tool iterations without producing a final answer]", False


# ---------------------------------------------------------------------------
# Pipeline
# ---------------------------------------------------------------------------
def _run_pipeline(query: str, model_id: Optional[str] = None,
                  runtime: Optional[dict] = None) -> dict:
    runtime = runtime or _sandbox_runtime()
    model = model_id or runtime.get("model") or DEFAULT_PROXY_MODEL
    guardrail_cfg = runtime.get("guardrail") or {"provider": "none"}
    agents = runtime.get("agents") or AGENTS
    steps: list[dict] = []
    user_ctx  = {"user_name": "triksha_sandbox", "user_email": "sandbox@triksha.internal"}

    dept      = _route_department(query, agents)
    sub_agent = _route_subagent(dept, query)

    steps += [
        {"step": "orchestrator_routing", "label": "Orchestrator — Department Routing",
         "department": dept["name"], "department_id": dept["id"], "description": dept["description"]},
        {"step": "subagent_routing", "label": f"{dept['name']} → Sub-agent Routing",
         "subagent": sub_agent["name"], "subagent_id": sub_agent["id"],
         "description": sub_agent["description"],
         "tools": [t["name"] for t in _tools_for(sub_agent)]},
    ]

    agent_ctx = {"agent_name": sub_agent["id"], "agent_id": sub_agent["id"]}

    # ── Inbound guardrail scan ─────────────────────────────────────────────
    inbound_data, inbound_timing = _guardrail_scan(
        [{"system": sub_agent["system_prompt"]}, {"query": query}],
        user_ctx, agent_ctx, phase="inbound", guardrail_cfg=guardrail_cfg,
    )
    steps.append({"step": "inbound_scan", "label": "Inbound Scan (User → Agent)",
                  "data": inbound_data, "guardrail_timing": inbound_timing})

    inbound_decision = None
    if isinstance(inbound_data, dict):
        res = inbound_data.get("result")
        if isinstance(res, dict):
            inbound_decision = res.get("decision")

    if inbound_decision == "block":
        steps.append({"step": "blocked", "agent": sub_agent["name"],
                      "text": "Request blocked by the guardrail — input flagged as a security threat."})
        return {"agent": sub_agent["name"], "department": dept["name"], "steps": steps}

    effective_query = query
    if isinstance(inbound_data, dict):
        res = inbound_data.get("result")
        if isinstance(res, dict):
            for m in (res.get("masked_content") or []):
                if isinstance(m, dict) and m.get("role") == "user":
                    effective_query = str(m.get("masked_text") or effective_query)
                    break

    # ── Agentic LLM loop (function calling) ────────────────────────────────
    steps.append({"step": "llm_call",
                  "label": f"Calling {sub_agent['name']} via LLM ({model}) — agentic mode",
                  "model": model,
                  "tools_available": [t["name"] for t in _tools_for(sub_agent)]})

    agent_response, llm_ok = _proxy_agentic(
        sub_agent, effective_query, steps, user_ctx, agent_ctx, model_id=model,
        guardrail_cfg=guardrail_cfg,
    )
    steps.append({"step": "llm_response", "text": agent_response, "llmOk": llm_ok})

    # ── Outbound guardrail scan ────────────────────────────────────────────
    outbound_data, outbound_timing = _guardrail_scan(
        [{"system": sub_agent["system_prompt"]},
         {"query": effective_query},
         {"response": agent_response}],
        user_ctx, agent_ctx, phase="outbound", guardrail_cfg=guardrail_cfg,
    )
    steps.append({"step": "outbound_scan", "label": "Outbound Scan (Agent → User)",
                  "data": outbound_data, "guardrail_timing": outbound_timing})

    final_response = agent_response
    if isinstance(outbound_data, dict):
        res = outbound_data.get("result")
        if isinstance(res, dict):
            od = res.get("decision")
            if od == "block":
                final_response = "Response blocked by the guardrail — agent output contained flagged content."
            elif od == "sanitize":
                for m in (res.get("masked_content") or []):
                    if isinstance(m, dict) and m.get("role") == "assistant":
                        final_response = str(m.get("masked_text") or final_response)
                        break

    steps.append({"step": "final_response", "agent": sub_agent["name"],
                  "department": dept["name"], "text": final_response})
    # Top-level `response` and `message` fields allow Triksha's agent scanner
    # (and any other client) to extract the answer without traversing steps[].
    return {
        "agent": sub_agent["name"],
        "department": dept["name"],
        "model": model,
        "response": final_response,
        "message": final_response,
        "steps": steps,
    }


# ---------------------------------------------------------------------------
# FastAPI endpoints
# ---------------------------------------------------------------------------
class ChatRequest(BaseModel):
    query: str
    # Optional proxy model override; validated against SUPPORTED_PROXY_MODELS.
    # Defaults to the DEFAULT_PROXY_MODEL env var if not provided.
    model: Optional[str] = None


@router.post("/sandbox/chat")
async def sandbox_chat(
    request: Request,
    body: ChatRequest,
):
    caller = _require_admin(request)
    if not body.query.strip():
        raise HTTPException(status_code=400, detail="query is required")

    runtime = _sandbox_runtime()
    # An explicit per-request model override (if a known proxy model) wins;
    # otherwise use the configured sandbox model.
    if body.model and body.model in SUPPORTED_PROXY_MODELS:
        model = body.model
    else:
        model = body.model or runtime.get("model") or DEFAULT_PROXY_MODEL

    import asyncio
    result = await asyncio.to_thread(_run_pipeline, body.query.strip(), model, runtime)

    ts = datetime.now(timezone.utc).isoformat()
    steps = result.get("steps", [])

    inbound_decision = ""
    outbound_decision = ""
    llm_ok = False
    final_response = ""
    for s in steps:
        step = s.get("step", "")
        if step == "inbound_scan":
            inbound_decision = (s.get("data") or {}).get("result", {}).get("decision", "") if isinstance((s.get("data") or {}), dict) else ""
        elif step == "outbound_scan":
            outbound_decision = (s.get("data") or {}).get("result", {}).get("decision", "") if isinstance((s.get("data") or {}), dict) else ""
        elif step == "llm_response":
            llm_ok = bool(s.get("llmOk"))
        elif step == "final_response":
            final_response = s.get("text", "")

    log_entry = {
        "ts": ts,
        "queried_by": caller,
        "query": body.query.strip(),
        "agent_name": result.get("agent", ""),
        "department": result.get("department", ""),
        "model": model,
        "inbound_decision": inbound_decision,
        "outbound_decision": outbound_decision,
        "llm_ok": llm_ok,
        "final_response": final_response,
        "steps": steps,
    }

    if _db is not None:
        try:
            _db.insert_sandbox_log(log_entry)
        except Exception as db_err:
            logger.warning("sandbox_log DB write failed: %s", db_err)

    _session_log.append({"ts": ts, "user": caller, "query": body.query.strip(), "result": result})
    if len(_session_log) > 200:
        _session_log.pop(0)

    return result


@router.get("/sandbox/health")
async def sandbox_health(request: Request):
    _require_admin(request)
    runtime = _sandbox_runtime()
    guardrail = runtime["guardrail"]
    provider = guardrail.get("provider", "none")
    model = runtime["model"]

    # Guardrail health: 'none' is always healthy; otherwise best-effort reachability.
    guardrail_ok = provider == "none"
    if provider in ("guardrail", "generic_http") and guardrail.get("base_url"):
        try:
            r = _req.get(f"{guardrail['base_url'].rstrip('/')}/health", timeout=5,
                         verify=guardrail.get("verify_ssl", False))
            guardrail_ok = r.status_code == 200
        except Exception:
            guardrail_ok = False
    elif provider == "connector":
        guardrail_ok = True  # configured; live check happens per-scan

    return {
        "guardrail": guardrail_ok,
        "guardrail_provider": provider,
        "guardrail_ok": guardrail_ok,
        "llm_ready": bool(_llm_api_key(model)),
        "model": model,
        "is_custom_agents": runtime["is_custom_agents"],
        "supported_models": sorted(SUPPORTED_PROXY_MODELS.keys()),
        "session_runs": len(_session_log),
    }


@router.get("/sandbox/config")
async def sandbox_config(request: Request):
    _require_admin(request)
    runtime = _sandbox_runtime()
    return {
        "guardrail_provider": runtime["guardrail"].get("provider", "none"),
        "model":              runtime["model"],
        "supported_models":   sorted(SUPPORTED_PROXY_MODELS.keys()),
        "llm_key_set":        bool(_llm_api_key(runtime["model"])),
        "is_custom_agents":   runtime["is_custom_agents"],
        "session_runs":       len(_session_log),
    }


# ── Customizable setup (instance-global, persisted via connectors_store) ───────
class SandboxSetup(BaseModel):
    guardrail: Optional[dict] = None   # {provider, base_url?, url?, protect_path?, verify_ssl?, connector_id?}
    guardrail_token: Optional[str] = None  # secret; blank = keep existing
    model: Optional[dict] = None       # {model}
    agents: Optional[list] = None      # custom agents; empty/None = default template


@router.get("/sandbox/setup")
async def get_sandbox_setup(request: Request):
    """Return the saved sandbox configuration (no secret values)."""
    _require_admin(request)
    try:
        import connectors_store
        stored = connectors_store.get_sandbox_config()
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not load setup: {exc}")
    cfg = stored.get("config", {}) or {}
    # Resolve guardrail connector options for the UI dropdown.
    guardrail_connectors = []
    try:
        import connectors_store as _cs
        for c in _cs.list_connectors():
            if c.get("type") in ("mcp",):
                continue  # only surface connectors usable as guardrails
            guardrail_connectors.append({"id": c["id"], "name": c["name"]})
    except Exception:
        pass
    return {
        "config": {
            "guardrail": cfg.get("guardrail") or {"provider": "default"},
            "model": cfg.get("model") or {"model": DEFAULT_PROXY_MODEL},
            "agents": cfg.get("agents") or [],
        },
        "secrets_set": stored.get("secrets_set", {}),
        "default_agents": AGENTS,
        "guardrail_providers": ["default", "none", "guardrail", "generic_http", "connector"],
        "guardrail_connectors": guardrail_connectors,
        "models": sorted(SUPPORTED_PROXY_MODELS.keys()),
    }


@router.put("/sandbox/setup")
async def put_sandbox_setup(request: Request, body: SandboxSetup):
    """Persist the sandbox configuration (instance-global)."""
    _require_admin(request)
    try:
        import connectors_store
        current = connectors_store.get_sandbox_config().get("config", {}) or {}
        new_cfg = dict(current)
        if body.guardrail is not None:
            new_cfg["guardrail"] = body.guardrail
        if body.model is not None:
            new_cfg["model"] = body.model
        if body.agents is not None:
            new_cfg["agents"] = body.agents
        secrets = {}
        if body.guardrail_token:
            secrets["guardrail_token"] = body.guardrail_token
        connectors_store.set_sandbox_config(new_cfg, secrets or None)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not save setup: {exc}")
    return {"status": "ok"}


@router.get("/sandbox/models")
async def sandbox_models(request: Request):
    """List supported models and which ones are currently usable."""
    _require_admin(request)
    return {
        "default": DEFAULT_PROXY_MODEL,
        "models": [
            {
                "id": model_id,
                "key_env": key_env,
                "available": bool(_llm_api_key(model_id)),
                "is_default": model_id == DEFAULT_PROXY_MODEL,
            }
            for model_id, key_env in SUPPORTED_PROXY_MODELS.items()
        ],
    }


@router.get("/sandbox/session")
async def sandbox_session(
    request: Request,
    limit: int = 200,
):
    _require_admin(request)
    # Prefer DB (persistent across restarts); fall back to in-memory list
    if _db is not None:
        try:
            rows = _db.get_sandbox_logs(limit=limit)
            runs = [
                {
                    "ts":    r.get("ts", ""),
                    "user":  r.get("queried_by", ""),
                    "query": r.get("query", ""),
                    "result": {
                        "agent":      r.get("agent_name", ""),
                        "department": r.get("department", ""),
                        "response":   r.get("final_response", ""),
                        "message":    r.get("final_response", ""),
                        "steps":      r.get("steps", []),
                    },
                }
                for r in rows
            ]
            return {"runs": runs, "total": len(runs), "source": "db"}
        except Exception as db_err:
            logger.warning("sandbox_session DB read failed, falling back to memory: %s", db_err)
    return {"runs": _session_log, "total": len(_session_log), "source": "memory"}


@router.delete("/sandbox/session")
async def sandbox_session_clear(request: Request):
    _require_admin(request)
    _session_log.clear()
    if _db is not None:
        try:
            _db.clear_sandbox_logs()
        except Exception as db_err:
            logger.warning("sandbox_session DB clear failed: %s", db_err)
    return {"cleared": True}


# ---------------------------------------------------------------------------
# DOCX report generation
# ---------------------------------------------------------------------------

def _decision_symbol(decision: str) -> str:
    d = (decision or "").lower()
    return {"allow": "✔ Allow", "pass": "✔ Pass", "block": "✘ Block",
            "sanitize": "⚠ Sanitize", "warn": "⚠ Warn",
            "service_not_available": "– Unavailable"}.get(d, decision or "–")


def _build_sandbox_docx(runs: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Base font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Title
    title = doc.add_heading("Triksha Sandbox — Security Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata table
    generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = [
        ("Generated", generated),
        ("Total Runs", str(len(runs))),
        ("Pipeline", "User → Guardrail Inbound → Orchestrator → Sub-agent → LLM → Guardrail Outbound → Response"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2, style="Light Grid Accent 1")
    for i, (k, v) in enumerate(meta):
        tbl.cell(i, 0).text = k
        tbl.cell(i, 1).text = v
        for run_para in tbl.cell(i, 0).paragraphs:
            for r in run_para.runs:
                r.bold = True
    doc.add_paragraph()

    if not runs:
        doc.add_paragraph("No runs recorded yet.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # Summary table
    doc.add_heading("Run Summary", level=1)
    hdr = ["#", "Time", "Query", "Agent", "Department", "Inbound", "Outbound"]
    summary_tbl = doc.add_table(rows=1 + len(runs), cols=len(hdr), style="Light Grid Accent 1")
    for j, h in enumerate(hdr):
        cell = summary_tbl.cell(0, j)
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, run in enumerate(runs):
        steps = run.get("result", {}).get("steps", [])
        in_dec = next((s.get("data", {}).get("result", {}).get("decision", "")
                       for s in steps if s.get("step") == "inbound_scan"), "")
        out_dec = next((s.get("data", {}).get("result", {}).get("decision", "")
                        for s in steps if s.get("step") == "outbound_scan"), "")
        ts_raw = run.get("ts", "")
        try:
            ts_str = datetime.fromisoformat(ts_raw.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M")
        except Exception:
            ts_str = ts_raw
        vals = [
            str(i + 1),
            ts_str,
            (run.get("query") or "")[:120],
            run.get("result", {}).get("agent", ""),
            run.get("result", {}).get("department", ""),
            _decision_symbol(in_dec),
            _decision_symbol(out_dec),
        ]
        for j, v in enumerate(vals):
            summary_tbl.cell(i + 1, j).text = v
    doc.add_paragraph()

    # Detailed run sections
    doc.add_heading("Run Details", level=1)
    for i, run in enumerate(runs):
        result = run.get("result", {})
        steps  = result.get("steps", [])
        ts_raw = run.get("ts", "")
        try:
            ts_str = datetime.fromisoformat(ts_raw.replace("Z", "+00:00")).strftime("%Y-%m-%d %H:%M UTC")
        except Exception:
            ts_str = ts_raw

        doc.add_heading(f"Run #{i + 1} — {ts_str}", level=2)

        # Query & routing
        p = doc.add_paragraph()
        p.add_run("Query: ").bold = True
        p.add_run(run.get("query", ""))

        p = doc.add_paragraph()
        p.add_run("Agent: ").bold = True
        p.add_run(f"{result.get('agent', '–')} · {result.get('department', '–')}")

        for step in steps:
            stype = step.get("step", "")

            if stype in ("inbound_scan", "outbound_scan"):
                label    = step.get("label", stype)
                data     = step.get("data") or {}
                res      = data.get("result") or {}
                decision = res.get("decision", "")
                timing   = step.get("guardrail_timing") or {}
                threats  = res.get("detections") or res.get("threats") or []
                masked   = res.get("masked_content") or []

                doc.add_heading(label, level=3)
                p = doc.add_paragraph()
                p.add_run("Decision: ").bold = True
                p.add_run(_decision_symbol(decision))

                p = doc.add_paragraph()
                p.add_run("Timing: ").bold = True
                p.add_run(
                    f"POST {timing.get('protect_post_ms', '–')}ms · "
                    f"Poll {timing.get('poll_attempts', 0)}×{timing.get('poll_http_round_trip_sum_ms', '–')}ms · "
                    f"Total {timing.get('total_round_trip_ms', '–')}ms · "
                    f"HTTP {timing.get('initial_http_status', '–')}"
                )

                if threats:
                    doc.add_paragraph("Threats detected:", style="List Bullet")
                    for t in threats:
                        conf = f" ({round(t.get('confidence', 0) * 100)}%)" if t.get("confidence") is not None else ""
                        doc.add_paragraph(
                            f"{t.get('label') or t.get('type') or str(t)}{conf}",
                            style="List Bullet 2",
                        )

                if masked:
                    doc.add_paragraph("Masked content:", style="List Bullet")
                    for m in masked:
                        doc.add_paragraph(
                            f"{m.get('role', '')}: {m.get('masked_text', '')}",
                            style="List Bullet 2",
                        )

            elif stype == "tool_call":
                doc.add_heading(step.get("label", "Tool Call"), level=3)
                p = doc.add_paragraph()
                p.add_run("Tool: ").bold = True
                p.add_run(step.get("tool_name", ""))
                p = doc.add_paragraph()
                p.add_run("Args: ").bold = True
                p.add_run(_json.dumps(step.get("tool_args", {})))
                grd_dec = step.get("guardrail_decision", "")
                if grd_dec:
                    p = doc.add_paragraph()
                    p.add_run("Guardrail scan: ").bold = True
                    p.add_run(_decision_symbol(grd_dec))

            elif stype == "llm_response" and step.get("text"):
                doc.add_heading("Raw LLM Output (before outbound scan)", level=3)
                doc.add_paragraph(step.get("text", ""))

            elif stype == "final_response":
                doc.add_heading("Final Response", level=3)
                doc.add_paragraph(step.get("text", ""))

        doc.add_paragraph()  # spacer between runs

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/sandbox/report.docx")
async def sandbox_report_docx(
    request: Request,
    limit: int = Query(default=200, le=1000),
):
    _require_admin(request)

    # Fetch runs from DB (preferred) or memory
    runs = []
    if _db is not None:
        try:
            rows = _db.get_sandbox_logs(limit=limit)
            runs = [
                {
                    "ts":    r.get("ts", ""),
                    "user":  r.get("queried_by", ""),
                    "query": r.get("query", ""),
                    "result": {
                        "agent":      r.get("agent_name", ""),
                        "department": r.get("department", ""),
                        "response":   r.get("final_response", ""),
                        "steps":      r.get("steps", []),
                    },
                }
                for r in rows
            ]
        except Exception as db_err:
            logger.warning("sandbox_report DB read failed, falling back to memory: %s", db_err)
            runs = _session_log
    else:
        runs = _session_log

    try:
        docx_bytes = await __import__("asyncio").get_event_loop().run_in_executor(
            None, _build_sandbox_docx, list(runs)
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx not installed")
    except Exception as e:
        logger.exception("DOCX generation failed: %s", e)
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")

    filename = f"sandbox-report-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S')}.docx"
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
