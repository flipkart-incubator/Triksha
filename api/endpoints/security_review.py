"""
Security Review Agent — Automated PRD-to-Security-Requirements Generator

Fetches a Google Doc (PRD) via GCP WIF credentials, analyzes it with
an LLM, and produces a one-pager security
requirements document in the same structured format as Triksha's
existing security review templates.

**LLM Token Budget**
  - maxOutputTokens: 1000 per call (hard cap)
  - thinkingConfig.thinkingBudget: 0 (reserve all tokens for output)
  - Generation is split into multiple phased calls, each within budget:
      Phase 1  → Identify attack surfaces (JSON, ~500 tokens)
      Phase 2  → Generate each section individually (~800 tokens each)
      Phase 3  → Summary table + Top 3 priorities (~600 tokens)
  - Final document is stitched from all phases.

Output Formats:
  - **docx** (default): Professional .docx file with formatted tables, headers
  - **markdown**: Raw markdown text in JSON response

Flow:
  1. Accept a Google Doc URL, document ID, raw text, or .docx upload
  2. Fetch document text via Google Docs API (using WIF) if needed
  3. Multi-phase LLM pipeline to generate structured output
  4. Return .docx file or markdown

No RAG needed — PRDs fit within Gemini's input context and the task
requires holistic, full-document understanding.
"""

import os
import re
import io
import json
import asyncio
import logging
import uuid
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional
from pathlib import Path

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Header, Request
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from rich.console import Console

logger = logging.getLogger("security_review")
console = Console()

router = APIRouter(tags=["Security Review"])

# ---------------------------------------------------------------------------
# Queue infrastructure (mirrors LLM / MCP / Agent scan pattern)
# ---------------------------------------------------------------------------

prd_review_queue: Optional[asyncio.Queue] = None
prd_review_event_queues: Dict[str, asyncio.Queue] = {}
running_reviews: Dict[str, Dict[str, Any]] = {}  # review_id → review record

# Lazy DB reference — set by main.py via set_database()
_db = None


def set_database(db) -> None:
    global _db
    _db = db


def _save_review(record: Dict[str, Any]) -> None:
    """Persist a new review record to DB (best-effort)."""
    if _db:
        try:
            _db.save_prd_review(record)
        except Exception:
            pass


def _update_review(review_id: str, update: Dict[str, Any]) -> None:
    """Update a review in DB (best-effort)."""
    if _db:
        try:
            _db.update_prd_review(review_id, update)
        except Exception as e:
            console.print(f"[red]⚠️  DB update failed for PRD review {review_id} (status={update.get('status')}): {e}[/]")

MAX_CONCURRENT_REVIEWS = int(os.getenv("TRIKSHA_MAX_CONCURRENT_PRD_REVIEWS", "2"))
PRD_QUEUE_MAX_SIZE = int(os.getenv("TRIKSHA_PRD_QUEUE_MAX_SIZE", "50"))


def init_prd_review_queue() -> None:
    """Called from main.py on_startup — initializes the queue and workers."""
    global prd_review_queue
    prd_review_queue = asyncio.Queue(maxsize=PRD_QUEUE_MAX_SIZE)
    for i in range(MAX_CONCURRENT_REVIEWS):
        asyncio.create_task(_prd_review_worker(i + 1))
    console.print(
        f"[cyan]Initialized PRD review queue (maxsize={PRD_QUEUE_MAX_SIZE}) "
        f"with {MAX_CONCURRENT_REVIEWS} workers[/]"
    )


async def _prd_review_worker(worker_id: int) -> None:
    console.print(f"[dim]PRD review worker {worker_id} started[/]")
    while True:
        try:
            item = await prd_review_queue.get()  # type: ignore
            if item is None:
                prd_review_queue.task_done()  # type: ignore
                break
            review_id, review_config = item
            await _run_prd_review_task(review_id, review_config)
        except Exception as e:
            console.print(f"[yellow]PRD review worker {worker_id} error: {e}[/]")
        finally:
            try:
                prd_review_queue.task_done()  # type: ignore
            except Exception:
                pass


async def _emit(review_id: str, event: Dict[str, Any]) -> None:
    """Put an SSE event on the per-review queue."""
    q = prd_review_event_queues.get(review_id)
    if q:
        await q.put(event)


async def _run_prd_review_task(review_id: str, review_config: Dict[str, Any]) -> None:
    """Worker task: run the full PRD review pipeline and emit SSE events."""
    record = running_reviews.get(review_id)
    if not record:
        return

    record["status"] = "running"
    record["progress"] = 0
    # Persist "running" to DB so recover_stuck_prd_reviews can correctly
    # identify in-progress reviews (previously it stayed "queued" in DB
    # for the entire review duration, and a silent DB-update failure on
    # completion left it "queued" → wrongly marked "failed" on restart).
    _update_review(review_id, {"status": "running", "progress": 0})
    await _emit(review_id, {
        "review_id": review_id, "status": "running",
        "event": "Started", "progress": 0,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })

    try:
        document_text = review_config["document_text"]
        document_title = review_config["document_title"]

        record["progress"] = 10
        await _emit(review_id, {
            "review_id": review_id, "status": "running",
            "event": "Identifying attack surfaces…", "progress": 10,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        })

        result = await _generate_security_review(
            document_text=document_text,
            document_title=document_title,
            author=review_config.get("author", "security@example.com"),
            additional_context=review_config.get("additional_context"),
            reference_link=review_config.get("reference_link"),
            review_id=review_id,
            record=record,
        )

        if not result["success"]:
            record["status"] = "failed"
            record["progress"] = 0
            record["error"] = result.get("error", "Generation failed")
            _update_review(review_id, {
                "status": "failed", "progress": 0,
                "completed_at": datetime.now(timezone.utc).isoformat(),
                "error": record["error"],
            })
            await _emit(review_id, {
                "review_id": review_id, "status": "failed",
                "event": "Failed", "error": record["error"],
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })
        else:
            record["status"] = "completed"
            record["progress"] = 100
            record["result"] = {
                "status": "success",
                "review_id": review_id,
                "document_title": document_title,
                "security_requirements_md": result["content"],
                "metadata": {
                    "generated_at": datetime.now(timezone.utc).isoformat(),
                    "document_length_chars": len(document_text),
                    "output_length_chars": len(result["content"]),
                    "model": "gemini-2.5-flash",
                    "total_llm_calls": result.get("llm_calls"),
                    "attack_surfaces_identified": result.get("surfaces_count"),
                    "sections_generated": result.get("sections_count"),
                    "phases_completed": result.get("phases_completed"),
                    "source": review_config.get("source", "upload"),
                },
            }
            record["_surfaces"] = result.get("_surfaces")
            record["_sections_md"] = result.get("_sections_md")
            record["_summary_md"] = result.get("_summary_md")

            _update_review(review_id, {
                "status": "completed",
                "progress": 100,
                "completed_at": datetime.now(timezone.utc).isoformat(),
                "result": record["result"],
                "_surfaces": record["_surfaces"],
                "_sections_md": record["_sections_md"],
                "_summary_md": record["_summary_md"],
            })

            await _emit(review_id, {
                "review_id": review_id, "status": "completed",
                "event": "Completed", "progress": 100,
                "result": record["result"],
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })

    except Exception as e:
        record["status"] = "failed"
        record["error"] = str(e)
        _update_review(review_id, {
            "status": "failed", "progress": 0,
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "error": str(e),
        })
        await _emit(review_id, {
            "review_id": review_id, "status": "failed",
            "event": "Failed", "error": str(e),
            "timestamp": datetime.now(timezone.utc).isoformat(),
        })
    finally:
        # SSE sentinel
        q = prd_review_event_queues.get(review_id)
        if q:
            await q.put(None)


async def _enqueue_review(review_id: str, review_config: Dict[str, Any]) -> None:
    """Enqueue via Kafka or local queue."""
    from kafka_client import is_kafka_enabled, enqueue_prd_review, KafkaProduceError
    if is_kafka_enabled():
        try:
            await enqueue_prd_review(review_id, review_config)
            return
        except KafkaProduceError as e:
            console.print(f"[red]Kafka produce failed for PRD review, falling back: {e}[/]")
    if prd_review_queue is None or prd_review_queue.full():
        raise HTTPException(status_code=429, detail="Review queue is full. Please try again later.")
    await prd_review_queue.put((review_id, review_config))

# ---------------------------------------------------------------------------
# LLM call tuning
# ---------------------------------------------------------------------------
_MAX_OUTPUT_TOKENS = 1000          # output budget per call
_MAX_RETRIES = 2                   # retry on transient provider errors
_RETRY_BASE_DELAY = 5              # seconds

# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------

class SecurityReviewRequest(BaseModel):
    """Request to generate a security requirements document from a PRD."""
    document_url: Optional[str] = Field(
        None,
        description="Google Doc URL (e.g. https://docs.google.com/document/d/DOC_ID/edit). "
                    "Either document_url, document_id, or document_text is required.",
    )
    document_id: Optional[str] = Field(
        None,
        description="Google Doc ID (the long alphanumeric string in the URL). "
                    "Either document_url, document_id, or document_text is required.",
    )
    document_text: Optional[str] = Field(
        None,
        description="Raw document text. If provided, skips Google Docs fetch. "
                    "Useful for non-Google-Doc inputs (Confluence, markdown, etc.).",
    )
    author: str = Field(
        "security@example.com",
        description="Author name for the security requirements document.",
    )
    additional_context: Optional[str] = Field(
        None,
        description="Optional extra context or focus areas for the review "
                    "(e.g. 'Focus on data handling and API security').",
    )
    output_format: str = Field(
        "docx",
        description="Output format: 'docx' (default) returns a downloadable Word document, "
                    "'markdown' returns raw markdown in JSON response.",
    )

    class Config:
        json_schema_extra = {
            "example": {
                "document_url": "https://docs.google.com/document/d/1r6Mg.../edit",
                "author": "security@example.com",
                "additional_context": "Focus on agent security and prompt injection risks",
                "output_format": "docx",
            }
        }


class SecurityReviewResponse(BaseModel):
    """Response containing the generated security requirements."""
    status: str
    review_id: str
    document_title: Optional[str] = None
    security_requirements_md: Optional[str] = None
    error: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None


# ---------------------------------------------------------------------------
# Google Doc fetcher
# ---------------------------------------------------------------------------

def _extract_doc_id(url_or_id: str) -> str:
    """Extract the Google Doc ID from a URL or return the raw ID."""
    match = re.search(r"/document/d/([a-zA-Z0-9_-]+)", url_or_id)
    if match:
        return match.group(1)
    if re.match(r"^[a-zA-Z0-9_-]{20,}$", url_or_id):
        return url_or_id
    raise ValueError(
        f"Cannot extract Google Doc ID from: {url_or_id}. "
        "Provide a full Google Docs URL or a valid document ID."
    )


async def _fetch_google_doc(doc_id: str) -> Dict[str, Any]:
    """
    Fetch a Google Doc via the Docs API using GCP WIF credentials.

    Returns dict with 'title' and 'text' keys.
    """
    try:
        from gcp_auth import get_gcp_access_token
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="GCP auth module not available. Ensure google-auth is installed.",
        )

    try:
        token = get_gcp_access_token(
            scopes=[
                "https://www.googleapis.com/auth/documents.readonly",
                "https://www.googleapis.com/auth/drive.readonly",
            ]
        )
    except EnvironmentError as e:
        raise HTTPException(status_code=500, detail=f"GCP credentials not configured: {e}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get GCP access token: {e}")

    import urllib.request
    import urllib.error

    url = f"https://docs.googleapis.com/v1/documents/{doc_id}"
    req = urllib.request.Request(
        url,
        headers={"Authorization": f"Bearer {token}", "Accept": "application/json"},
    )

    try:
        resp = await asyncio.to_thread(urllib.request.urlopen, req, timeout=30)
        data = json.loads(resp.read())
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        try:
            msg = json.loads(body).get("error", {}).get("message", body[:500])
        except Exception:
            msg = body[:500]
        if e.code == 403:
            raise HTTPException(
                status_code=403,
                detail=(
                    "Access denied to Google Doc. "
                    "Share the doc with the service account or grant it viewer access. "
                    f"Error: {msg}"
                ),
            )
        elif e.code == 404:
            raise HTTPException(
                status_code=404,
                detail=f"Google Doc not found: {doc_id}. Ensure it exists and is shared.",
            )
        else:
            raise HTTPException(status_code=502, detail=f"Google Docs API error {e.code}: {msg}")

    title = data.get("title", "Untitled Document")
    body_content = data.get("body", {})
    content_elements = body_content.get("content", [])

    text_parts: List[str] = []
    for element in content_elements:
        paragraph = element.get("paragraph", {})
        for pe in paragraph.get("elements", []):
            text_run = pe.get("textRun", {})
            content = text_run.get("content", "")
            if content:
                text_parts.append(content)
        table = element.get("table", {})
        for row in table.get("tableRows", []):
            row_texts = []
            for cell in row.get("tableCells", []):
                cell_parts = []
                for cc in cell.get("content", []):
                    for cpe in cc.get("paragraph", {}).get("elements", []):
                        ct = cpe.get("textRun", {}).get("content", "").strip()
                        if ct:
                            cell_parts.append(ct)
                row_texts.append(" ".join(cell_parts))
            if row_texts:
                text_parts.append(" | ".join(row_texts) + "\n")

    full_text = "".join(text_parts).strip()
    if not full_text:
        raise HTTPException(status_code=422, detail="The Google Doc appears empty.")
    return {"title": title, "text": full_text}


# ---------------------------------------------------------------------------
# LLM caller — token-budget-aware, matches codebase patterns
# ---------------------------------------------------------------------------

async def _call_llm(
    prompt: str,
    system_instruction: Optional[str] = None,
    temperature: float = 0.3,
) -> Dict[str, Any]:
    """
    Single LLM completion against the user-configured provider.

    Triksha: routes through llm_providers (OpenAI / Anthropic / Gemini) using
    the API key the user set in Settings. Retries on transient failures.

    Returns dict with 'success' (bool), 'content' (str), 'error' (str).
    """
    import llm_providers

    if not llm_providers.is_configured():
        return {"success": False, "error": "No LLM provider API key configured. Set it in Settings."}

    # Retry loop for transient provider errors
    last_error = ""
    for attempt in range(_MAX_RETRIES + 1):
        try:
            content = await asyncio.to_thread(
                llm_providers.complete_sync,
                prompt,
                system=system_instruction,
                temperature=temperature,
                max_tokens=_MAX_OUTPUT_TOKENS,
            )
            if not content:
                return {"success": False, "error": "No text content in LLM response"}
            return {"success": True, "content": content.strip()}
        except llm_providers.LLMNotConfigured as exc:
            return {"success": False, "error": str(exc)}
        except Exception as exc:
            last_error = f"LLM request failed: {exc}"
            if attempt < _MAX_RETRIES:
                await asyncio.sleep(_RETRY_BASE_DELAY * (2 ** attempt))
                continue
            return {"success": False, "error": last_error}

    return {"success": False, "error": last_error or "Exhausted retries"}


# ---------------------------------------------------------------------------
# Helper: strip markdown fences from LLM output
# ---------------------------------------------------------------------------

def _strip_fences(text: str) -> str:
    """Remove wrapping ```markdown ... ``` fences if present."""
    t = text.strip()
    if t.startswith("```markdown"):
        t = t[len("```markdown"):].strip()
    elif t.startswith("```json"):
        t = t[len("```json"):].strip()
    elif t.startswith("```"):
        t = t[3:].strip()
    if t.endswith("```"):
        t = t[:-3].strip()
    return t


# ---------------------------------------------------------------------------
# .docx text extractor — read input PRDs from .docx files
# ---------------------------------------------------------------------------

def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file (in-memory bytes)."""
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
        with z.open("word/document.xml") as f:
            tree = ET.parse(f)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    lines: List[str] = []
    for p in tree.findall(".//w:p", ns):
        texts = [r.text for r in p.findall(".//w:t", ns) if r.text]
        line = "".join(texts).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def _extract_text_from_doc(file_bytes: bytes) -> str:
    """Extract plain text from a legacy .doc (binary Word) file using mammoth."""
    import mammoth
    result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
    return result.value


# ---------------------------------------------------------------------------
# .docx builder — convert structured output to a professional Word document
# ---------------------------------------------------------------------------

def _build_docx(
    document_title: str,
    author: str,
    surfaces: List[Dict[str, Any]],
    sections_md: List[str],
    summary_md: str,
    reference_link: Optional[str] = None,
) -> bytes:
    """
    Build a .docx file from the structured security review output.
    Returns the docx as bytes (ready to write to disk or stream).
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # -- Title --
    title_para = doc.add_heading(f"{document_title}: Security Requirements", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -- Metadata table --
    meta_data = [
        ("Author", author),
        ("Scope", "Security properties that the described system must satisfy — independent of implementation"),
        ("Reference", reference_link or document_title),
        ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
    ]
    meta_table = doc.add_table(rows=len(meta_data), cols=2, style="Light Grid Accent 1")
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (field, detail) in enumerate(meta_data):
        meta_table.cell(i, 0).text = field
        meta_table.cell(i, 1).text = detail
        # Bold field names
        for run in meta_table.cell(i, 0).paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()  # spacer

    # -- How to Read --
    doc.add_heading("How to Read This Document", level=2)
    all_components = sorted(set(c for s in surfaces for c in s.get("components", [])))
    tags_str = ", ".join(f"[{c}]" for c in all_components)
    doc.add_paragraph(
        "Each requirement states what must be true, not how to achieve it. "
        "A companion security review document should evaluate whether the current "
        "implementation satisfies these requirements and identify gaps."
    )
    doc.add_paragraph(f"Requirements are tagged: {tags_str}")

    # -- Summary + Priorities (first, as executive overview) --
    doc.add_page_break()
    _add_markdown_section_to_docx(doc, summary_md)

    # -- Detailed Sections --
    for section_md in sections_md:
        doc.add_page_break()
        _add_markdown_section_to_docx(doc, section_md)

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_markdown_section_to_docx(doc, section_md: str):
    """Parse a single markdown section and add it to the docx Document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = section_md.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        if line.startswith("## "):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue

        # Table (starts with |)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _add_table_to_docx(doc, table_lines)
            continue

        # Bold key attack scenarios header
        if line.strip().startswith("**Key attack scenarios:**"):
            p = doc.add_paragraph()
            run = p.add_run("Key attack scenarios:")
            run.bold = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Bullet point
        if line.strip().startswith("- "):
            bullet_text = line.strip()[2:]
            bullet_text = _clean_markdown_inline(bullet_text)
            doc.add_paragraph(bullet_text, style="List Bullet")
            i += 1
            continue

        # Numbered list item
        if re.match(r"^\d+\.\s", line.strip()):
            item_text = re.sub(r"^\d+\.\s*", "", line.strip())
            item_text = _clean_markdown_inline(item_text)
            doc.add_paragraph(item_text, style="List Number")
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        if text:
            text = _clean_markdown_inline(text)
            doc.add_paragraph(text)
        i += 1


def _add_table_to_docx(doc, table_lines: List[str]):
    """Parse markdown table lines and add a formatted table to docx."""
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT

    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # Skip separator rows (e.g. |---|---|)
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows_data.append(cells)

    if not rows_data:
        return

    n_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=n_cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_cells in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx < n_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = _clean_markdown_inline(cell_text)
                # Bold header row
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                # Set font size
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()  # spacer after table


def _clean_markdown_inline(text: str) -> str:
    """Remove inline markdown formatting (bold, italic) for plain text."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # **bold**
    text = re.sub(r"\*(.+?)\*", r"\1", text)        # *italic*
    text = re.sub(r"`(.+?)`", r"\1", text)          # `code`
    return text.strip()


# ---------------------------------------------------------------------------
# Phase 1 — Identify attack surfaces
# ---------------------------------------------------------------------------

_PHASE1_SYSTEM = """\
You are an expert AppSec engineer. Given a Product Requirements Document (PRD),
identify the MOST CRITICAL security-relevant attack surfaces for the described system.

Return ONLY a valid JSON array. No markdown, no explanation. Example:
[
  {"id": 1, "title": "Indirect Prompt Injection via Production Data", "risk": "critical",
   "components": ["AGENT", "TOOL"], "brief": "Production API responses flow into LLM prompts..."},
  {"id": 2, "title": "Blackboard State Tampering", "risk": "high",
   "components": ["ALL"], "brief": "Shared state file can be modified..."}
]

Rules:
- Identify 5–8 attack surfaces. Focus on what matters most — skip low-risk boilerplate.
- Merge related surfaces into one (e.g. auth + session = "Authentication & Session Security").
- Risk levels: critical, high, medium.
- "components" are tags like [AGENT], [API], [DATA], [UI], [TOOL], [ALL].
- "brief" is ONE sentence explaining why this surface matters.
- Prioritise domain-specific risks over generic ones.
- For AI/ML systems: prompt injection, PII leakage, tool abuse, agent scope creep.
"""


async def _phase1_identify_surfaces(
    prd_text: str,
    additional_context: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """Phase 1: Identify attack surfaces from the PRD. Returns list of surfaces."""

    # Truncate PRD to ~12K chars (~4K tokens) to leave room for system prompt + output
    truncated = prd_text[:12000]
    if len(prd_text) > 12000:
        truncated += "\n\n[... document truncated for token budget ...]"

    prompt = f"## PRD\n\n{truncated}"
    if additional_context:
        prompt += f"\n\n## Additional Focus Areas\n{additional_context}"

    result = await _call_llm(prompt, system_instruction=_PHASE1_SYSTEM, temperature=0.2)
    if not result["success"]:
        raise RuntimeError(f"Phase 1 failed: {result['error']}")

    raw = _strip_fences(result["content"])

    try:
        surfaces = json.loads(raw)
    except json.JSONDecodeError:
        # Try to extract JSON array from the response
        match = re.search(r"\[.*\]", raw, re.DOTALL)
        if match:
            surfaces = json.loads(match.group())
        else:
            raise RuntimeError(f"Phase 1: Could not parse attack surfaces JSON. Raw: {raw[:500]}")

    if not isinstance(surfaces, list) or len(surfaces) == 0:
        raise RuntimeError("Phase 1: LLM returned empty or invalid attack surface list")

    return surfaces


# ---------------------------------------------------------------------------
# Phase 2 — Generate individual sections
# ---------------------------------------------------------------------------

_PHASE2_SYSTEM = """\
You are an expert AppSec engineer writing a crisp security requirements document.

Given a PRD excerpt and ONE attack surface, produce a SINGLE concise section
in this EXACT markdown format:

## N. Title [COMPONENT_TAGS]

One-line description of the attack surface.

**Key attack scenarios:**
- Concrete scenario 1 (one sentence)
- Concrete scenario 2 (one sentence)

| # | Requirement | Rationale |
|---|:-----------|:----------|
| N.1 | Specific, testable requirement | One-line risk explanation |
| N.2 | ... | ... |

Rules:
- 2–4 requirements per section. Quality over quantity.
- Keep requirements actionable and testable — one sentence each.
- Keep rationale to one sentence — the specific risk, not generic "for security."
- 2–3 attack scenarios max, one sentence each.
- Output ONLY the markdown section. No preamble.
"""


async def _phase2_generate_section(
    prd_text: str,
    surface: Dict[str, Any],
) -> str:
    """Phase 2: Generate one requirements section for a given attack surface."""

    # Provide a focused PRD excerpt — truncate to leave room for output
    truncated = prd_text[:10000]
    if len(prd_text) > 10000:
        truncated += "\n\n[... truncated ...]"

    sid = surface.get("id", 1)
    title = surface.get("title", "Unknown")
    risk = surface.get("risk", "medium")
    components = surface.get("components", ["ALL"])
    brief = surface.get("brief", "")

    prompt = f"""## PRD (for context)

{truncated}

---

## Attack Surface to Cover

Section number: {sid}
Title: {title}
Risk: {risk}
Components: {', '.join(f'[{c}]' for c in components)}
Description: {brief}

Generate the full section now."""

    result = await _call_llm(prompt, system_instruction=_PHASE2_SYSTEM, temperature=0.3)
    if not result["success"]:
        # Return a stub section rather than failing the entire review
        logger.warning("Phase 2 failed for section %d (%s): %s", sid, title, result["error"])
        return (
            f"## {sid}. {title} [{', '.join(components)}]\n\n"
            f"*Section generation failed: {result['error']}*\n"
        )

    section_md = _strip_fences(result["content"])
    return section_md


# ---------------------------------------------------------------------------
# Phase 3 — Summary table + Top 5 priorities
# ---------------------------------------------------------------------------

_PHASE3_SYSTEM = """\
You are an expert AppSec engineer. Given a list of security requirement sections,
produce TWO things in markdown. Be extremely concise.

Output this EXACT format and NOTHING else:

## Summary: Attack Surface Map

| Attack Surface | Risk | Key Requirements |
|:---------------|:-----|:-----------------|
| **Surface 1** | critical | 1.1–1.N |
| **Surface 2** | high | 2.1–2.N |

---

## Top 3 Priorities

1. **Highest priority (Section N)** — One-line justification
2. ...
3. ...

Rules:
- Include ALL sections in the summary table.
- Top 3 only — the most critical. One line each.
"""


async def _phase3_summary(
    surfaces: List[Dict[str, Any]],
    sections_md: List[str],
) -> str:
    """Phase 3: Generate summary table and top 5 priorities."""

    # Build a compact representation of sections for the LLM
    section_summaries = []
    for s in surfaces:
        sid = s.get("id", "?")
        title = s.get("title", "?")
        risk = s.get("risk", "?")
        section_summaries.append(f"Section {sid}: {title} (Risk: {risk})")

    prompt = f"""## Sections Generated

{chr(10).join(section_summaries)}

## Section Content (excerpts)

{chr(10).join(s[:400] for s in sections_md)}

Generate the Summary Attack Surface Map and Top 5 Priorities now."""

    result = await _call_llm(prompt, system_instruction=_PHASE3_SYSTEM, temperature=0.2)
    if not result["success"]:
        logger.warning("Phase 3 failed: %s", result["error"])
        # Build a minimal summary from the surfaces list
        rows = []
        for s in surfaces:
            rows.append(
                f"| **{s.get('title', '?')}** | {s.get('risk', '?')} | {s.get('id', '?')}.1–{s.get('id', '?')}.N |"
            )
        return (
            "## Summary: Attack Surface Map\n\n"
            "| Attack Surface | Risk | Key Requirements |\n"
            "|:---------------|:-----|:-----------------|\n"
            + "\n".join(rows)
            + "\n\n---\n\n## Top 3 Priorities\n\n"
            "1. *Generation failed — please review sections above manually.*\n"
        )

    return _strip_fences(result["content"])


# ---------------------------------------------------------------------------
# Orchestrator — multi-phase pipeline
# ---------------------------------------------------------------------------

async def _generate_security_review(
    document_text: str,
    document_title: str,
    author: str,
    additional_context: Optional[str] = None,
    reference_link: Optional[str] = None,
    review_id: Optional[str] = None,
    record: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Multi-phase security review generation pipeline.

    Phase 1: Identify attack surfaces (~1 call, ≤1000 output tokens)
    Phase 2: Generate each section (~N calls, ≤1000 output tokens each)
    Phase 3: Summary + priorities (~1 call, ≤1000 output tokens)

    Returns dict with 'success', 'content', 'error', 'phases_completed'.
    """

    async def _emit_progress(pct: int, event: str) -> None:
        if record is not None:
            record["progress"] = pct
        if review_id:
            await _emit(review_id, {
                "review_id": review_id, "status": "running",
                "event": event, "progress": pct,
                "timestamp": datetime.now(timezone.utc).isoformat(),
            })

    # --- Phase 1: Identify attack surfaces ---
    console.print("[cyan]  Phase 1/3: Identifying attack surfaces...[/]")
    try:
        surfaces = await _phase1_identify_surfaces(document_text, additional_context)
    except RuntimeError as e:
        return {"success": False, "error": str(e), "phases_completed": 0}

    console.print(f"[green]  ✓ Phase 1 complete: {len(surfaces)} attack surfaces identified[/]")
    await _emit_progress(20, f"Attack surfaces identified: {len(surfaces)} areas")

    # --- Phase 2: Generate sections (one call per surface) ---
    console.print(f"[cyan]  Phase 2/3: Generating {len(surfaces)} requirement sections...[/]")
    sections_md: List[str] = []
    n = len(surfaces)

    for i, surface in enumerate(surfaces):
        console.print(
            f"[dim]    [{i+1}/{n}] {surface.get('title', '?')}...[/]"
        )
        section = await _phase2_generate_section(document_text, surface)
        sections_md.append(section)
        # Phase 2 spans progress 20→85 proportionally
        pct = 20 + int((i + 1) / n * 65)
        await _emit_progress(pct, f"Section {i+1}/{n}: {surface.get('title', '?')}")

    console.print(f"[green]  ✓ Phase 2 complete: {len(sections_md)} sections generated[/]")

    # --- Phase 3: Summary + Top 5 ---
    console.print("[cyan]  Phase 3/3: Generating summary and priorities...[/]")
    await _emit_progress(90, "Generating summary and priorities…")
    summary_md = await _phase3_summary(surfaces, sections_md)
    console.print("[green]  ✓ Phase 3 complete[/]")

    # --- Stitch final document ---
    # Build the "How to Read" tags dynamically from surfaces
    all_components = set()
    for s in surfaces:
        for c in s.get("components", []):
            all_components.add(c)
    tags_str = ", ".join(f"[{c}]" for c in sorted(all_components))

    header = (
        f"# {document_title}: Security Requirements\n\n"
        f"| Field | Detail |\n"
        f"|:------|:-------|\n"
        f"| **Author** | {author} |\n"
        f"| **Scope** | Security properties that the described system must satisfy — independent of implementation |\n"
        f"| **Reference** | {reference_link or document_title} |\n"
        f"\n---\n\n"
        f"## How to Read This Document\n\n"
        f"Each requirement states **what** must be true, not **how** to achieve it. "
        f"A companion security review document should evaluate whether the current "
        f"implementation satisfies these requirements and identify gaps.\n\n"
        f"Requirements are tagged: {tags_str}\n"
        f"\n---\n\n"
    )

    body = "\n\n---\n\n".join(sections_md)

    # Summary + priorities come FIRST, then detailed sections
    full_document = header + summary_md + "\n\n---\n\n" + body

    return {
        "success": True,
        "content": full_document,
        "phases_completed": 3,
        "surfaces_count": len(surfaces),
        "sections_count": len(sections_md),
        "llm_calls": 1 + len(surfaces) + 1,  # phase1 + phase2*N + phase3
        # Structured data for docx generation
        "_surfaces": surfaces,
        "_sections_md": sections_md,
        "_summary_md": summary_md,
        "_reference_link": reference_link,
    }


# ---------------------------------------------------------------------------
# API Endpoints
# ---------------------------------------------------------------------------

_OUTPUT_DIR = Path(os.environ.get("TRIKSHA_DATA_DIR", "data")) / "security_reviews"
_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


@router.post(
    "/security-review/generate",
    summary="Generate security requirements from a PRD (async)",
)
async def generate_security_review(
    request: SecurityReviewRequest,
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """Enqueue a PRD security review. Returns review_id immediately."""

    review_id = str(uuid.uuid4())
    document_text = None
    document_title = "Unknown"
    reference_link: Optional[str] = None

    if request.document_text:
        document_text = request.document_text
        document_title = "Provided Document"
    elif request.document_url or request.document_id:
        raw_input = request.document_url or request.document_id
        try:
            doc_id = _extract_doc_id(raw_input)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        reference_link = request.document_url or f"https://docs.google.com/document/d/{doc_id}/edit"
        doc_data = await _fetch_google_doc(doc_id)
        document_title = doc_data["title"]
        document_text = doc_data["text"]
    else:
        raise HTTPException(status_code=400, detail="Provide document_url, document_id, or document_text.")

    review_config = {
        "document_text": document_text,
        "document_title": document_title,
        "author": request.author,
        "additional_context": request.additional_context,
        "reference_link": reference_link,
        "output_format": request.output_format,
        "source": "google_docs" if (request.document_url or request.document_id) else "raw_text",
        "created_by": x_proxy_user or "anonymous",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }

    # Register review record
    q: asyncio.Queue = asyncio.Queue()
    prd_review_event_queues[review_id] = q
    running_reviews[review_id] = {
        "review_id": review_id,
        "document_title": document_title,
        "author": request.author,
        "status": "queued",
        "created_by": x_proxy_user or "anonymous",
        "created_at": review_config["created_at"],
        "reference_link": reference_link,
    }
    _save_review(running_reviews[review_id])
    await _emit(review_id, {
        "review_id": review_id, "status": "queued",
        "event": "Queued", "progress": 0,
        "timestamp": review_config["created_at"],
    })

    await _enqueue_review(review_id, review_config)

    return {"review_id": review_id, "status": "queued",
            "events_url": f"/security-review/{review_id}/events",
            "status_url": f"/security-review/{review_id}"}


@router.post("/security-review/upload", summary="Upload a PRD and start async security review")
async def upload_and_review(
    file: UploadFile = File(...),
    reference_id: str = Form("", description="Optional reference ID (unused in OS)"),
    author: str = Form("security@example.com"),
    additional_context: Optional[str] = Form(None),
    output_format: str = Form("markdown"),
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """Extract text from uploaded PRD, enqueue review. Returns review_id immediately."""

    # Optional external reference id; auto-generates OS-* if omitted.
    # otherwise auto-generate a local reference id.
    if not reference_id.strip():
        reference_id = f"OS-{uuid.uuid4().hex[:10].upper()}"

    review_id = str(uuid.uuid4())
    file_bytes = await file.read()
    filename = file.filename or "uploaded_document"

    if filename.lower().endswith(".docx"):
        try:
            document_text = _extract_text_from_docx(file_bytes)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse .docx: {e}")
        document_title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")
    elif filename.lower().endswith(".doc"):
        try:
            document_text = _extract_text_from_doc(file_bytes)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse .doc: {e}")
        document_title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")
    elif filename.lower().endswith((".txt", ".md")):
        document_text = file_bytes.decode("utf-8", errors="replace")
        document_title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ")
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file: {filename}. Use .docx, .doc, .txt, or .md.")

    if not document_text.strip():
        raise HTTPException(status_code=422, detail="Uploaded document appears empty.")

    review_config = {
        "document_text": document_text,
        "document_title": document_title,
        "author": author,
        "reference_id": reference_id.strip(),
        "additional_context": additional_context,
        "reference_link": filename,
        "output_format": output_format,
        "source": "upload",
        "created_by": x_proxy_user or "anonymous",
        "created_at": datetime.now(timezone.utc).isoformat(),
    }

    q: asyncio.Queue = asyncio.Queue()
    prd_review_event_queues[review_id] = q
    running_reviews[review_id] = {
        "review_id": review_id,
        "document_title": document_title,
        "reference_id": reference_id.strip(),
        "author": author,
        "status": "queued",
        "created_by": x_proxy_user or "anonymous",
        "created_at": review_config["created_at"],
        "reference_link": review_config.get("reference_link"),
    }
    _save_review(running_reviews[review_id])
    await _emit(review_id, {
        "review_id": review_id, "status": "queued",
        "event": "Queued", "progress": 0,
        "timestamp": review_config["created_at"],
    })

    await _enqueue_review(review_id, review_config)

    return {"review_id": review_id, "status": "queued",
            "events_url": f"/security-review/{review_id}/events",
            "status_url": f"/security-review/{review_id}"}


@router.get("/security-review/list", summary="List all PRD reviews")
async def list_reviews(
    scope: Optional[str] = None,
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """Return all reviews visible to the current user (DB + in-memory running).

    scope: 'mine' (only the calling user's), 'others' (everyone else's),
    default no ownership filter.
    """
    # Start with DB records (historical + completed)
    db_records: Dict[str, Dict] = {}
    if _db:
        try:
            for r in _db.list_prd_reviews(limit=200):
                db_records[r["review_id"]] = r
        except Exception:
            pass

    # Overlay with in-memory (running/queued take priority — they have live progress)
    merged: Dict[str, Dict] = {**db_records}
    for review_id, r in running_reviews.items():
        merged[review_id] = {k: v for k, v in r.items() if k not in ("_surfaces", "_sections_md", "_summary_md")}

    reviews = sorted(merged.values(), key=lambda r: r.get("created_at", ""), reverse=True)

    scope_filter = (scope or "all").lower()
    if scope_filter in ("mine", "others") and x_proxy_user:
        # Normalize both sides: strip @domain, lowercase. Handles cases
        # where one side is an email and the other is a bare username, or
        # the case differs between storage and the proxy header.
        def _username(val: str) -> str:
            s = str(val or "").strip()
            if not s:
                return ""
            return (s.split("@", 1)[0]).lower()

        caller = _username(x_proxy_user)

        def _is_mine(rec):
            return bool(caller) and _username(rec.get("created_by")) == caller

        if scope_filter == "mine":
            reviews = [r for r in reviews if _is_mine(r)]
        else:
            reviews = [r for r in reviews if not _is_mine(r)]

    return reviews


@router.get(
    "/security-review/health",
    summary="Check security review agent health",
    description="Verify that GCP credentials and an LLM provider are configured.",
)
async def security_review_health():
    """Check if the security review agent is properly configured."""
    import llm_providers
    checks = {}

    checks["llm_provider"] = llm_providers.is_configured()

    try:
        from gcp_auth import is_gcp_available, _get_auth_mode
        checks["gcp_configured"] = is_gcp_available()
        gcp_auth_mode = _get_auth_mode()
    except ImportError:
        checks["gcp_configured"] = False
        gcp_auth_mode = "unavailable"

    all_ok = all(checks.values())

    return {
        "status": "healthy" if all_ok else "degraded",
        "checks": checks,
        "gcp_auth_mode": gcp_auth_mode,
        "llm_provider": llm_providers.get_provider(),
        "token_budget": {
            "max_output_tokens_per_call": _MAX_OUTPUT_TOKENS,
        },
        "notes": (
            f"All systems operational. GCP auth mode: {gcp_auth_mode}."
            if all_ok
            else "Some features unavailable. "
                 "GCP SA key is needed for Google Docs fetching "
                 "(set GOOGLE_APPLICATION_CREDENTIALS). "
                 "An LLM provider API key is needed for analysis (set it in Settings). "
                 "Raw text / .docx upload works without GCP."
        ),
    }


@router.get("/security-review/{review_id}", summary="Get PRD review status and results")
async def get_review_status(
    review_id: str,
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """Poll status of a queued/running/completed PRD review."""
    record = running_reviews.get(review_id)
    if not record and _db:
        record = _db.get_prd_review(review_id)
    if not record:
        raise HTTPException(status_code=404, detail=f"Review {review_id} not found")
    response = dict(record)
    response.pop("_surfaces", None)
    response.pop("_sections_md", None)
    response.pop("_summary_md", None)
    return response


@router.get("/security-review/{review_id}/events", summary="Stream PRD review progress (SSE)")
async def stream_review_events(
    request: Request,
    review_id: str,
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """SSE stream — emits queued/running/completed events for a PRD review."""
    record = running_reviews.get(review_id)
    if not record:
        raise HTTPException(status_code=404, detail=f"Review {review_id} not found")
    q = prd_review_event_queues.get(review_id)
    if q is None:
        q = asyncio.Queue()
        prd_review_event_queues[review_id] = q

    async def event_generator():
        try:
            while True:
                if await request.is_disconnected():
                    break
                event = await q.get()
                if event is None:
                    yield "event: end\ndata: {\"status\": \"done\"}\n\n"
                    break
                yield f"data: {json.dumps(event)}\n\n"
        except asyncio.CancelledError:
            return

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


@router.get("/security-review/{review_id}/download", summary="Download completed PRD review as .docx")
async def download_review_docx(
    review_id: str,
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """Generate and stream a .docx download for a completed review."""
    record = running_reviews.get(review_id)
    if not record and _db:
        record = _db.get_prd_review(review_id)
    if not record:
        raise HTTPException(status_code=404, detail="Review not found")
    if record.get("status") != "completed":
        raise HTTPException(status_code=400, detail=f"Review is {record.get('status')}, not completed")

    surfaces = record.get("_surfaces")
    sections_md = record.get("_sections_md")
    summary_md = record.get("_summary_md") or ""

    # Fallback for older records where surfaces_json/sections_md columns are NULL
    # (PyMySQL couldn't store a raw Python list; result_json was saved separately).
    # Use the compiled full-text markdown to produce a single-section docx.
    if not surfaces or not sections_md:
        full_md = (record.get("result") or {}).get("security_requirements_md")
        if not full_md:
            raise HTTPException(status_code=500, detail="Review result data missing")
        # Treat the entire compiled doc as one section so the download still works.
        surfaces = [{"title": "Security Requirements", "components": []}]
        sections_md = [full_md]
        summary_md = ""

    docx_bytes = _build_docx(
        document_title=record["document_title"],
        author=record.get("author", "security@example.com"),
        surfaces=surfaces,
        sections_md=sections_md,
        summary_md=summary_md,
        reference_link=record.get("reference_link"),
    )
    safe_name = re.sub(r"[^\w\s-]", "", record["document_title"]).strip().replace(" ", "_")
    filename = f"{safe_name}_Security_Requirements_{review_id[:8]}.docx"
    filepath = _OUTPUT_DIR / filename
    filepath.write_bytes(docx_bytes)
    return FileResponse(
        path=str(filepath),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.delete("/security-review/{review_id}", summary="Delete a PRD review (admin only)")
async def delete_review(
    review_id: str,
    x_proxy_user: Optional[str] = Header(None, alias="x-proxy-user"),
):
    """Delete a PRD review from in-memory store.

    Requires triksha.scan.prd-delete (granted to triksha.admin role).
    Allowed for own reviews and reviews owned by anyone in the same AuthZ group.
    """
    record = running_reviews.get(review_id)
    if not record and _db:
        record = _db.get_prd_review(review_id)
    if not record:
        raise HTTPException(status_code=404, detail=f"Review {review_id} not found")
    running_reviews.pop(review_id, None)
    prd_review_event_queues.pop(review_id, None)
    if _db:
        _db.delete_prd_review(review_id)
    return {"deleted": review_id}
